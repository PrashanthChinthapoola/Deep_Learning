# document_processor.py

import streamlit as st
import docx
from io import BytesIO
# Import PyMuPDF
import fitz # PyMuPDF is imported as fitz

from langchain_openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document as LangchainDocument

# Assume embeddings_model is initialized in config and accessible
# from config import embeddings_model # Or pass it as argument - still passing embeddings

@st.cache_data
def extract_text_from_docx(docx_file_content):
    """Cached function to read text from bytes content of a .docx file."""
    if st.session_state.get("show_tool_messages"):
         st.info(f"Extracting text from DOCX...")

    try:
        doc = docx.Document(BytesIO(docx_file_content))
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        if st.session_state.get("show_tool_messages"):
             st.success(f"Text extraction complete.")
        return text
    except Exception as e:
        st.error(f"Error reading DOCX file: {e}")
        return None

@st.cache_data # Use st.cache_data for text extraction
def extract_text_from_pdf(pdf_file_content):
    """Cached function to read text from bytes content of a .pdf file using PyMuPDF."""
    if st.session_state.get("show_tool_messages"):
        st.info(f"Extracting text from PDF...")

    text = ""
    try:
        # Open the PDF from bytes
        pdf_document = fitz.open(stream=pdf_file_content, filetype="pdf")
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            text += page.get_text() + "\n" # Add a newline between pages
        pdf_document.close()

        if st.session_state.get("show_tool_messages"):
            st.success(f"Text extraction complete.")

        if not text.strip():
             st.warning("PDF text extraction resulted in empty content.")
             return None # Return None if no text was extracted

        return text
    except Exception as e:
        st.error(f"Error reading PDF file: {e}")
        return None


@st.cache_resource
def create_vector_store(text_content, _embeddings_model):
    """
    Cached function to create a FAISS vector store.
    Takes embeddings_model as argument, now prefixed with _ for caching.
    """
    if st.session_state.get("show_tool_messages"):
         st.info(f"Creating vector store for document search...")
    try:
        # Check if text_content is valid before splitting
        if not text_content or len(text_content.strip()) < 100:
            st.warning(f"Document content is too short or empty. Cannot create vector store.")
            return None

        docs = [LangchainDocument(page_content=text_content)]
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=200)
        split_docs = text_splitter.split_documents(docs)

        if not split_docs:
            st.warning(f"Document could not be split into chunks. It might be too short or unstructured.")
            return None

        # Use the parameter name with the underscore inside the function
        vector_store = FAISS.from_documents(split_docs, _embeddings_model)
        if st.session_state.get("show_tool_messages"):
             st.success(f"Vector store created with {len(split_docs)} chunks.")
        return vector_store.as_retriever()
    except Exception as e:
        st.error(f"Error creating vector store: {e}")
        if "RateLimitError" in str(e):
             st.error(f"You might be hitting API rate limits. Try again later or check your plan.")
        elif "embedding model" in str(e).lower() or "model not found" in str(e).lower():
             st.error(f"The specified OpenAI embedding model might not be available or your account doesn't have access. Try 'text-embedding-ada-002' or check your OpenAI dashboard.")
        return None