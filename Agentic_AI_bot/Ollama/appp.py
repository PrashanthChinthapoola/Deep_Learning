import streamlit as st
import openai
import json
import os
from dotenv import load_dotenv
import docx
from langchain_openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document as LangchainDocument
import time
from datetime import datetime
# import uuid

# --- Configuration and Setup ---
load_dotenv()
openai_api_key = os.environ.get("OPENAI_API_KEY")

if not openai_api_key:
    st.error("OpenAI API key not found in .env file.")
    st.markdown("Please create a `.env` file in the same directory as your script with `OPENAI_API_KEY='your_key_here'`.")
    st.stop()

try:
    client = openai.OpenAI(api_key=openai_api_key)
    embeddings_model = OpenAIEmbeddings(openai_api_key=openai_api_key, model="text-embedding-ada-002")
except Exception as e:
    st.error(f"Failed to initialize OpenAI client or embeddings model: {e}")
    st.markdown("Please double-check your OpenAI API key, your internet connection, and if your account has access to the embedding model.")
    st.stop()

st.set_page_config(page_title="Agentic Doc Chatbot + Time", layout="wide")
st.title("📄⏳ Agentic Chatbot for Your DOCX Document + Time")
st.caption("Upload a DOCX file in the sidebar to chat about it within the selected session. I can also tell you the current time and answer general knowledge questions.")


# --- Session State Initialization for MULTIPLE SESSIONS ---
if "sessions" not in st.session_state:
    st.session_state.sessions = {}
    st.session_state.session_counter = 0
    first_session_id = f"Chat {st.session_state.session_counter + 1}"
    st.session_state.sessions[first_session_id] = {
        "messages": [{"role": "assistant", "content": "Hello! Upload a DOCX in the sidebar to chat about it, or ask about the current time or general knowledge."}],
        "doc_retriever": None,
        "uploaded_file_name": None,
        "name": first_session_id
    }
    st.session_state.current_session_id = first_session_id
    st.session_state.session_counter += 1

if "show_tool_messages" not in st.session_state:
     st.session_state.show_tool_messages = False

# --- Helper Functions ---
@st.cache_data
def extract_text_from_docx(docx_file_content):
    """Cached function to read text from bytes content of a .docx file."""
    if st.session_state.get("show_tool_messages"):
        st.info("Extracting text from DOCX...")
    try:
        from io import BytesIO
        doc = docx.Document(BytesIO(docx_file_content))
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        if st.session_state.get("show_tool_messages"):
             st.success("Text extraction complete.")
        return text
    except Exception as e:
        st.error(f"Error reading DOCX file: {e}")
        return None

@st.cache_resource
def create_vector_store(text_content, _embeddings_model):
    """
    Cached function to create a FAISS vector store.
    _embeddings_model is prefixed with _ to tell Streamlit not to hash this unhashable object.
    """
    if not text_content or len(text_content.strip()) < 100:
        st.warning("Document content is too short or empty. Cannot create vector store.")
        return None
    if st.session_state.get("show_tool_messages"):
        st.info("Creating vector store for document search...")
    try:
        docs = [LangchainDocument(page_content=text_content)]
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=200)
        split_docs = text_splitter.split_documents(docs)

        if not split_docs:
            st.warning("Document could not be split into chunks. It might be too short or unstructured.")
            return None

        vector_store = FAISS.from_documents(split_docs, _embeddings_model)
        if st.session_state.get("show_tool_messages"):
             st.success(f"Vector store created with {len(split_docs)} chunks.")
        return vector_store.as_retriever()
    except Exception as e:
        st.error(f"Error creating vector store: {e}")
        if "RateLimitError" in str(e):
            st.error("You might be hitting API rate limits. Try again later or check your plan.")
        elif "embedding model" in str(e).lower() or "model not found" in str(e).lower():
            st.error("The specified OpenAI embedding model might not be available or your account doesn't have access. Try 'text-embedding-ada-002' or check your OpenAI dashboard.")
        return None

# --- Tool Implementations ---
def search_document(query: str) -> str:
    """
    Searches the uploaded DOCX document using the vector store of the CURRENT SESSION.
    Returns relevant document chunks as a string.
    Returns a specific error message if the vector store is not initialized for the current session.
    """
    retriever = st.session_state.sessions[st.session_state.current_session_id].get("doc_retriever")

    if st.session_state.get("show_tool_messages"):
        st.info(f"🛠️ Calling tool: `search_document` with query: '{query}'")

    if not retriever:
        return "TOOL_ERROR: Document vector store not initialized for this session. Please upload a document in the sidebar."
    try:
        time.sleep(0.5) # Simulate tool latency
        relevant_docs = retriever.invoke(query)

        if not relevant_docs:
            if st.session_state.get("show_tool_messages"):
                st.info("✔️ `search_document` found no relevant sections.")
            return "TOOL_RESPONSE: No relevant information found in the document for that query in this session."

        context = "\n\n---\n\n".join([doc.page_content for doc in relevant_docs])
        if st.session_state.get("show_tool_messages"):
            st.info(f"✔️ `search_document` found {len(relevant_docs)} relevant section(s).")
        max_context_length = 5000 # Prevent sending excessively long context to the model
        if len(context) > max_context_length:
            context = context[:max_context_length] + "\n\n[... context truncated ...]"
            if st.session_state.get("show_tool_messages"):
                st.warning(f"Context from document truncated to {max_context_length} characters.")

        return "TOOL_RESPONSE: Relevant document sections found:\n\n" + context

    except Exception as e:
        st.error(f"Error during document search tool execution: {e}")
        return f"TOOL_ERROR: An error occurred while searching the document: {e}"

def summarize_section(text: str) -> str:
    """Summarizes a given text section."""
    if not text or len(text.strip()) < 50:
        if st.session_state.get("show_tool_messages"):
            st.warning("Input text for summarization is too short.")
        return "TOOL_RESPONSE: Cannot summarize a very short or empty text section."

    if st.session_state.get("show_tool_messages"):
        st.info("🛠️ Calling tool: `summarize_section`...")
    try:
        time.sleep(0.5) # Simulate tool latency
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful assistant skilled at summarizing text concisely and accurately."},
                {"role": "user", "content": f"Please summarize the following text:\n\n{text}"},
            ],
            max_tokens=300,
            temperature=0.7
        )
        summary = response.choices[0].message.content
        if st.session_state.get("show_tool_messages"):
             st.info("✔️ `summarize_section` complete.")
        return "TOOL_RESPONSE: Summary:\n\n" + summary
    except Exception as e:
        st.error(f"Error during summarization tool execution: {e}")
        return f"TOOL_ERROR: An error occurred while trying to summarize the section: {e}"

def get_current_datetime_tool() -> str:
    """Gets the current date and time."""
    if st.session_state.get("show_tool_messages"):
        st.info("🛠️ Calling tool: `get_current_datetime`...")
    try:
        now = datetime.now()
        formatted_datetime = now.strftime("%Y-%m-%d %H:%M:%S %A")
        if st.session_state.get("show_tool_messages"):
             st.info(f"✔️ `get_current_datetime` complete: {formatted_datetime}")
        return f"TOOL_RESPONSE: Current datetime is {formatted_datetime}"
    except Exception as e:
        st.error(f"Error getting current datetime: {e}")
        return f"TOOL_ERROR: An error occurred while getting the current date and time: {e}"


# --- Tool Definitions for OpenAI API ---
tools = [
    {
        "type": "function",
        "function": {
            "name": "search_document",
            "description": "Searches the uploaded DOCX document for information within the currently active chat session. **THIS IS THE PRIMARY TOOL FOR ANSWERING QUESTIONS ABOUT THE DOCUMENT'S CONTENT.** Use this tool whenever the user's question asks about or implies needing information from the document, even if the relevance is not perfectly clear. Formulate the best possible search query based on the user's request. You MUST handle the specific response prefixes 'TOOL_ERROR:' and 'TOOL_RESPONSE:' properly.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "The search query or keywords to search within the document."
                    }
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "summarize_section",
            "description": "Summarizes a block of text that has been retrieved from the document using `search_document`. Use this when you have retrieved a long section of text and need to provide a concise overview, or if the user explicitly asks for a summary of something found in the document. Handle 'TOOL_RESPONSE:' and 'TOOL_ERROR:'.",
            "parameters": {
                "type": "object",
                "properties": {
                    "text": {
                        "type": "string",
                        "description": "The text content from the document that needs to be summarized."
                    }
                },
                "required": ["text"],
            },
        },
    },
    {
       "type": "function",
       "function": {
            "name": "get_current_datetime",
            "description": "Gets the current date and time. Use this tool whenever the user explicitly asks for the current time or date. Handle 'TOOL_RESPONSE:' and 'TOOL_ERROR:'.",
            "parameters": {
                "type": "object",
                "properties": {},
            },
       },
    }
]

# Map function names to the actual Python functions
available_functions = {
    "search_document": search_document,
    "summarize_section": summarize_section,
    "get_current_datetime": get_current_datetime_tool
}


# --- Sidebar ---
with st.sidebar:
    # Add a main title for the sidebar for overall context
    st.title("Chat Options")

    st.subheader("Chat Sessions")

    # Display buttons for each session
    session_ids = list(st.session_state.sessions.keys())
    # Optional: Sort sessions, maybe by creation time if stored, or alphabetically
    # session_ids.sort()

    for session_id in session_ids:
        is_current = (session_id == st.session_state.current_session_id)
        display_name = st.session_state.sessions[session_id].get("name", session_id)
        # Use Markdown to bold and add an arrow for the current session
        button_label = f"➡️ **{display_name}**" if is_current else display_name

        # Use a unique key for each button
        # Using the session ID as part of the key is a good practice
        if st.button(button_label, key=f"session_button_{session_id}", use_container_width=True):
            if not is_current:
                st.session_state.current_session_id = session_id
                st.rerun() # Trigger rerun to load the selected session

    # Button to create a new session - place it logically after the session list
    if st.button("➕ Start New Chat", use_container_width=True):
        st.session_state.session_counter += 1
        new_session_id = f"Chat {st.session_state.session_counter}"
        while new_session_id in st.session_state.sessions:
             st.session_state.session_counter += 1
             new_session_id = f"Chat {st.session_state.session_counter}"

        st.session_state.sessions[new_session_id] = {
            "messages": [{"role": "assistant", "content": "Hello! Upload a DOCX in the sidebar to chat about it, or ask about the current time or general knowledge."}],
            "doc_retriever": None,
            "uploaded_file_name": None,
            "name": new_session_id
        }
        st.session_state.current_session_id = new_session_id
        st.rerun()

    st.markdown("---") # Separator

    # Clear Current Chat Button - place it near the session controls it affects
    # Added spacing above it for visual separation from the new chat button
    st.write("") # Add some space
    if st.button(f"🗑️ Clear '{st.session_state.current_session_id}' Chat & Doc", use_container_width=True):
         # Add a confirmation step? (Optional but good UX for destructive action)
         # For simplicity, direct clear for now
         st.session_state.sessions[st.session_state.current_session_id] = {
             "messages": [{"role": "assistant", "content": "Chat and document cleared for this session. Upload a new DOCX to get started."}],
             "doc_retriever": None,
             "uploaded_file_name": None,
             "name": st.session_state.current_session_id
         }
         st.rerun()


    st.markdown("---") # Separator

    st.subheader("Upload Document")
    # The file uploader state is tied to the currently selected session
    current_session_state = st.session_state.sessions[st.session_state.current_session_id]
    current_session_uploaded_file_name = current_session_state["uploaded_file_name"]

    uploaded_file = st.file_uploader(
        f"Choose DOCX for '{st.session_state.current_session_id}'",
        type=["docx"],
        help="Select a .docx file to process for the currently selected chat session.",
        key=f"uploader_session_{st.session_state.current_session_id}" # Unique key per session
    )

    st.markdown("---") # Separator

    # Use an expander for settings to keep them collapsed by default
    with st.expander("Settings"):
        st.subheader("Display Settings")
        # This setting applies globally, not per session
        st.session_state.show_tool_messages = st.checkbox(
            "Show Agent/Tool Steps",
            value=st.session_state.show_tool_messages,
            key="show_tool_messages_checkbox",
            help="Show intermediate steps where the agent calls tools."
        )

    st.markdown("---") # Separator

    # Use another expander for the About section
    with st.expander("About"):
        st.write("This is an agentic chatbot that can process a DOCX document (per session) and answer questions based on its content using function calling.")
        st.write("It can also answer general questions and provide the current time.")
        st.write("It uses Streamlit for the interface and OpenAI's GPT models with function calling.")
        st.write("Powered by OpenAI and Streamlit.")
        # Optional: Add link to source code, author, etc.

    # Optional: Add something at the very bottom of the sidebar
    # st.markdown("---")
    # st.write("App Version 1.0")


# --- File Processing Logic ---
if uploaded_file is not None:
    uploaded_file_content = uploaded_file.getvalue()
    uploaded_file_name = uploaded_file.name

    current_session_state = st.session_state.sessions[st.session_state.current_session_id]
    if current_session_state["uploaded_file_name"] != uploaded_file_name or current_session_state["doc_retriever"] is None:

        if st.session_state.show_tool_messages:
             st.session_state.sessions[st.session_state.current_session_id]["messages"].append({"role": "assistant", "content": f"Processing '{uploaded_file_name}' for '{st.session_state.current_session_id}'..."})

        with st.spinner(f"Processing '{uploaded_file_name}' for '{st.session_state.current_session_id}'..."):
            doc_text = extract_text_from_docx(uploaded_file_content)

            if doc_text:
                retriever = create_vector_store(doc_text, embeddings_model)

                st.session_state.sessions[st.session_state.current_session_id]["doc_retriever"] = retriever
                st.session_state.sessions[st.session_state.current_session_id]["uploaded_file_name"] = uploaded_file_name

                if retriever:
                    if st.session_state.show_tool_messages:
                         st.session_state.sessions[st.session_state.current_session_id]["messages"].append({"role": "assistant", "content": f"✅ '{uploaded_file_name}' processed successfully for '{st.session_state.current_session_id}'! You can now ask questions about its content."})
                    else:
                         st.session_state.sessions[st.session_state.current_session_id]["messages"].append({"role": "assistant", "content": f"✅ '{uploaded_file_name}' processed for '{st.session_state.current_session_id}'!"})
                else:
                    st.session_state.sessions[st.session_state.current_session_id]["messages"].append({"role": "assistant", "content": f"⚠️ Failed to create a searchable index for '{uploaded_file_name}'. The document might be empty or have formatting issues."})
            else:
                st.session_state.sessions[st.session_state.current_session_id]["messages"].append({"role": "assistant", "content": f"⚠️ Failed to read text from '{uploaded_file_name}'. Please ensure it's a valid .docx file."})

        st.rerun()


# --- Display Chat Messages ---
current_messages = st.session_state.sessions[st.session_state.current_session_id]["messages"]

for message in current_messages:
    if message["role"] != "tool":
        with st.chat_message(message["role"]):
            if message.get("content") is not None:
                if st.session_state.show_tool_messages and message["role"] == "assistant" and message.get("tool_calls"):
                    st.markdown("🤖 **Thinking... (Tool Call)**")
                st.markdown(message["content"])
            elif st.session_state.show_tool_messages and message["role"] == "assistant" and message.get("tool_calls"):
                 st.markdown("🤖 **Thinking... (Tool Call)**")


# --- Chat Input ---
prompt = st.chat_input(f"Ask about the document ({st.session_state.sessions[st.session_state.current_session_id]['uploaded_file_name'] or 'No document uploaded'}), current time, or general knowledge in '{st.session_state.current_session_id}'...")


# --- Define System Prompt ---
system_prompt = """You are a helpful and versatile AI assistant with access to specialized tools and general knowledge. Your goal is to provide accurate and relevant information.

**Absolutely Critical Instruction:** Whenever the user asks a question that **could possibly be answered by or relate to the content of the uploaded document in the CURRENT CHAT SESSION**, you **MUST call the `search_document` tool first** to attempt to find the answer within the document. Do this even if you think you might know the answer from your general knowledge. The document is your primary source for any document-like query when available in the current session.

Here are your capabilities and strict instructions:

1.  **Document Questions:** If the user's query is about, refers to, or **even potentially relates to the document in the current session**:
    * **IMMEDIATELY call the `search_document` tool.**
    * **Process the Tool's Response (which will start with "TOOL_RESPONSE:" or "TOOL_ERROR:"):**
        * If `search_document` returns "TOOL_ERROR: Document vector store not initialized...", **you MUST respond to the user** by explaining clearly that no document has been uploaded and processed *for this chat session* yet and they need to upload one in the sidebar to ask document-specific questions *for this session*. Do NOT try to answer document questions from general knowledge in this case.
        * If `search_document` returns "TOOL_RESPONSE: No relevant information found...", **you MUST respond to the user** by stating clearly that the information was not found *in the uploaded document for this session*.
        * If `search_document` returns "TOOL_RESPONSE: Relevant document sections found:\n\n...", carefully analyze the provided document text. **Synthesize your answer SOLELY based on this text.** Do NOT use your general knowledge or invent information.
        * If `search_document` returns any other "TOOL_ERROR:", inform the user there was an error searching the document.
    * **Use `summarize_section`:** If `search_document` provided text and a summary is needed or requested, you *may* then call `summarize_section` on that text and incorporate the summary into your final answer. Handle its "TOOL_RESPONSE:" and "TOOL_ERROR:" results.

2.  **Current Time Question:** If the user explicitly asks for the current time or date:
    * Call the `get_current_datetime` tool.
    * Process its "TOOL_RESPONSE:" or "TOOL_ERROR:" and provide the information to the user.

3.  **General Knowledge Questions:** If the user asks a simple factual question that is **clearly and undeniably NOT** related to the document *in the current session* or the current time (e.g., "What is the capital of France?", "Explain photosynthesis.", "Who invented the telephone?"):
    * Answer the question directly using your built-in general knowledge.
    * **DO NOT call any tools** for these types of questions.

**Response Formulation:**
* Your final answer to the user should be clear, helpful, and directly address their query.
* **Do not show the "TOOL_RESPONSE:" or "TOOL_ERROR:" prefixes to the user.** These are for your internal processing only.
* Clearly indicate if your answer came from the document (e.g., "According to the document...", "The current time is...").

**In summary: Document-like query -> Use `search_document` -> Act strictly on its output (including the error message for missing document for the CURRENT SESSION). Time query -> Use `get_current_datetime` -> Act on its output. Clearly Non-document/Non-time query -> Answer directly.**
"""


# --- Main Chat Processing Logic ---
if prompt:
    current_messages = st.session_state.sessions[st.session_state.current_session_id]["messages"]
    current_messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    messages_for_api = []
    messages_for_api.append({"role": "system", "content": system_prompt})

    for m in current_messages:
        if m["role"] == "system":
            continue
        msg = {"role": m["role"], "content": m.get("content", "")}
        if m["role"] == "assistant" and m.get("tool_calls"):
             msg["tool_calls"] = m["tool_calls"]
        elif m["role"] == "tool" and m.get("tool_call_id"):
             msg["tool_call_id"] = m["tool_call_id"]
             msg["name"] = m["name"]
             msg["content"] = str(m.get("content", ""))
        messages_for_api.append(msg)

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=messages_for_api,
            tools=tools,
            tool_choice="auto",
            temperature=0.7
        )

        response_message = response.choices[0].message
        tool_calls = response_message.tool_calls

        if tool_calls:
            st.session_state.sessions[st.session_state.current_session_id]["messages"].append(response_message.model_dump(exclude_unset=True))
            if st.session_state.show_tool_messages:
                 st.info("🤖 Model decided to use tool(s)...")

            messages_after_tool_calls = messages_for_api + [response_message.model_dump(exclude_unset=True)]

            for tool_call in tool_calls:
                function_name = tool_call.function.name
                function_to_call = available_functions.get(function_name)

                if function_to_call:
                    try:
                        function_args = json.loads(tool_call.function.arguments)
                        with st.spinner(f"Executing tool: `{function_name}`..." if st.session_state.show_tool_messages else "..."):
                             function_response_content = function_to_call(**function_args)

                        tool_output_message = {
                            "tool_call_id": tool_call.id,
                            "role": "tool",
                            "name": function_name,
                            "content": function_response_content,
                        }
                        st.session_state.sessions[st.session_state.current_session_id]["messages"].append(tool_output_message)
                        messages_after_tool_calls.append(tool_output_message)

                        if st.session_state.show_tool_messages:
                             st.success(f"Tool `{function_name}` executed.")

                    except json.JSONDecodeError:
                        error_msg_content = f"TOOL_ERROR: Invalid JSON arguments provided by model for tool `{function_name}`. Args: {tool_call.function.arguments}"
                        st.error(error_msg_content)
                        tool_output_message = {
                             "tool_call_id": tool_call.id,
                             "role": "tool",
                             "name": function_name,
                             "content": error_msg_content,
                        }
                        st.session_state.sessions[st.session_state.current_session_id]["messages"].append(tool_output_message)
                        messages_after_tool_calls.append(tool_output_message)

                    except Exception as e:
                        error_msg_content = f"TOOL_ERROR: An error occurred during execution of `{function_name}`: {e}"
                        st.error(error_msg_content)
                        tool_output_message = {
                            "tool_call_id": tool_call.id,
                            "role": "tool",
                            "name": function_name,
                            "content": error_msg_content,
                        }
                        st.session_state.sessions[st.session_state.current_session_id]["messages"].append(tool_output_message)
                        messages_after_tool_calls.append(tool_output_message)

                else:
                    error_msg_content = f"TOOL_ERROR: Model attempted to call an unimplemented or unknown tool: {function_name}"
                    st.error(error_msg_content)
                    tool_output_message = {
                         "tool_call_id": tool_call.id,
                         "role": "tool",
                         "name": function_name,
                         "content": error_msg_content,
                    }
                    st.session_state.sessions[st.session_state.current_session_id]["messages"].append(tool_output_message)
                    messages_after_tool_calls.append(tool_output_message)

            if st.session_state.show_tool_messages:
                st.info("🧠 Getting final response based on tool results...")
            try:
                final_response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=messages_after_tool_calls,
                )
                final_answer = final_response.choices[0].message.content
                st.session_state.sessions[st.session_state.current_session_id]["messages"].append({"role": "assistant", "content": final_answer})
                with st.chat_message("assistant"):
                    st.markdown(final_answer)

            except Exception as e:
                st.error(f"An error occurred during the second API call (synthesizing response): {e}")
                error_message = f"Sorry, I encountered an error while trying to formulate my response based on the tool results: {e}"
                st.session_state.sessions[st.session_state.current_session_id]["messages"].append({"role": "assistant", "content": error_message})
                with st.chat_message("assistant"):
                    st.markdown("Sorry, I encountered an error while trying to formulate my response.")

        else:
            final_answer = response_message.content
            st.session_state.sessions[st.session_state.current_session_id]["messages"].append({"role": "assistant", "content": final_answer})
            with st.chat_message("assistant"):
                st.markdown(final_answer)

    except openai.AuthenticationError as e:
        st.error(f"Authentication Error: {e}. Please check your OpenAI API key.")
        error_message = "Authentication failed. Please check your OpenAI API key."
        st.session_state.sessions[st.session_state.current_session_id]["messages"].append({"role": "assistant", "content": error_message})
        with st.chat_message("assistant"):
            st.markdown("Authentication failed. Please check your OpenAI API key.")

    except Exception as e:
        st.error(f"An error occurred during the API call: {e}")
        error_message = f"An error occurred during the conversation: {e}"
        st.session_state.sessions[st.session_state.current_session_id]["messages"].append({"role": "assistant", "content": error_message})
        with st.chat_message("assistant"):
            st.markdown("Sorry, an unexpected error occurred.")